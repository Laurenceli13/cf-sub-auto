"""Generate VPN subscription import manual as Word document with embedded QR codes."""

import io
from pathlib import Path
import qrcode
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

TOKEN = "9eeeb06655487c286527001b"
HOST = "sub.tamovpn.top"
SUB_URL = f"https://{HOST}/sub.txt?token={TOKEN}"
CLASH_URL = f"https://{HOST}/sub.txt?token={TOKEN}&format=clash"
WEB_URL = f"https://{HOST}/"

# jsDelivr CDN mirrors for China accessibility
JS_BASE = "https://cdn.jsdelivr.net/gh/Laurenceli13/cf-sub-auto@main/public"
SUB_MIRROR_URL = f"{JS_BASE}/sub_v2ray.txt"
CLASH_MIRROR_URL = f"{JS_BASE}/sub_clash.yaml"

# GitHub mirror proxy for client downloads
GHPROXY = "https://ghproxy.com/https://github.com"

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "VPN订阅导入操作手册.docx"
QR_DIR = ROOT / "public" / "qrcodes"
QR_DIR.mkdir(parents=True, exist_ok=True)


def make_qr(url: str, filename: str) -> io.BytesIO:
    """Generate QR code PNG in memory and save to disk."""
    img = qrcode.make(url, border=2)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    path = QR_DIR / filename
    img.save(path)
    return buf


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers: list[str], rows: list[list[str]], col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")
    return table


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =========================================
    # COVER PAGE
    # =========================================
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VPN 订阅导入操作手册")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("跨境电商网络工具 · 内部使用")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        f"订阅域名: {HOST}\n"
        f"更新时间: 每6小时自动刷新\n"
        f"版本: v2.0"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x47, 0x56, 0x69)

    for _ in range(6):
        doc.add_paragraph("")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("机密文件 · 请勿外传")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)

    doc.add_page_break()

    # =========================================
    # SECTION 1: QR CODE SCAN
    # =========================================
    doc.add_heading("一、扫码导入（推荐）", level=1)

    doc.add_paragraph(
        "使用手机客户端扫描下方二维码，即可自动导入订阅。"
        "无需手动输入任何地址。"
    )

    # Generate QR codes
    qr_sub = make_qr(SUB_URL, "sub_qr.png")
    qr_clash = make_qr(CLASH_URL, "clash_qr.png")

    # QR code table: side by side
    qr_table = doc.add_table(rows=2, cols=2)
    qr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: QR images
    for col_idx, (qr_buf, label) in enumerate([
        (qr_sub, "V2Ray 订阅"),
        (qr_clash, "Clash 订阅")
    ]):
        cell = qr_table.rows[0].cells[col_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(qr_buf, width=Inches(2.0))

    # Row 1: Labels
    for col_idx, label in enumerate([
        "V2Ray 通用订阅",
        "Clash Meta 订阅"
    ]):
        cell = qr_table.rows[1].cells[col_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)

    doc.add_paragraph("")

    # Subscription info box
    info_box = doc.add_table(rows=4, cols=1)
    info_box.style = 'Light Shading Accent 1'
    info_box.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        f"订阅地址: {SUB_URL}",
        f"Clash 地址: {CLASH_URL}",
        f"Web 控制台: {WEB_URL}",
        f"Token: {TOKEN}"
    ]
    for i, text in enumerate(info_data):
        cell = info_box.rows[i].cells[0]
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Consolas'

    doc.add_paragraph("")

    h2_note = doc.add_paragraph()
    run = h2_note.add_run(
        "⚠️ 注意：请勿将订阅地址和 Token 泄露给公司外部人员"
    )
    run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
    run.bold = True
    run.font.size = Pt(10)

    doc.add_page_break()

    # =========================================
    # SECTION 2: v2rayNG (Android)
    # =========================================
    doc.add_heading("二、v2rayNG（Android）配置教程", level=1)

    doc.add_paragraph(
        "v2rayNG 是 Android 平台最常用的 V2Ray 客户端，"
        "支持 vmess/vless/trojan/ss 等协议。"
    )

    steps_v2ray = [
        ("1", "下载安装",
         "1. 从 GitHub 下载最新版 v2rayNG:\n"
         "   https://github.com/2dust/v2rayNG/releases\n"
         "2. 下载 .apk 文件并安装（需允许「未知来源」安装）。"),
        ("2", "扫码导入（推荐）",
         "1. 打开 v2rayNG\n"
         "2. 右上角「+」→ 「扫描二维码」\n"
         "3. 扫描本手册第一页的「V2Ray 通用订阅」二维码即可。"),
        ("3", "手动导入（备选）",
         "1. 打开 v2rayNG\n"
         "2. 左上角「三」菜单 → 「订阅设置」\n"
         "3. 右上角「+」\n"
         "4. 别名填「公司VPN」\n"
         "5. URL 填入:\n"
         f"   {SUB_URL}\n"
         "6. 点击「确定」\n"
         "7. 返回主页 → 右上角菜单 → 「更新订阅」"),
        ("4", "连接使用",
         "1. 更新完成后主界面出现节点列表\n"
         "2. 选择延迟最低的节点\n"
         "3. 点击右下角「V」图标连接\n"
         "4. 首次使用需允许 VPN 权限"),
    ]

    for num, title, desc in steps_v2ray:
        doc.add_heading(f"步骤{num}: {title}", level=3)
        for line in desc.split("\n"):
            if line.strip():
                doc.add_paragraph(line, style='List Bullet')

    doc.add_page_break()

    # =========================================
    # SECTION 3: Shadowrocket (iOS)
    # =========================================
    doc.add_heading(
        "三、Shadowrocket / 小火箭（iOS）配置教程",
        level=1
    )

    doc.add_paragraph(
        "Shadowrocket（小火箭）是 iOS 平台最流行的代理工具。\n"
        "需要外区 Apple ID 下载（美区/港区等），售价约 $2.99。"
    )

    steps_sr = [
        ("1", "下载安装",
         "1. 使用外区 Apple ID 登录 App Store\n"
         "2. 搜索「Shadowrocket」\n"
         "3. 购买并下载"),
        ("2", "扫码导入（推荐）",
         "1. 打开 Shadowrocket\n"
         "2. 右上角「+」\n"
         "3. 类型选「Subscribe」\n"
         "4. 点击 URL 输入框右侧的「扫描」图标\n"
         "5. 扫描本手册第一页的「V2Ray 通用订阅」二维码"),
        ("3", "手动导入（备选）",
         "1. 打开 Shadowrocket\n"
         "2. 右上角「+」\n"
         "3. 类型选「Subscribe」\n"
         "4. URL 输入框填入:\n"
         f"   {SUB_URL}\n"
         "5. 备注填「公司VPN」\n"
         "6. 右上角「完成」"),
        ("4", "更新与连接",
         "1. 主界面下拉刷新订阅\n"
         "2. 选择一个节点\n"
         "3. 打开顶部开关连接\n"
         "4. 首次使用需允许 VPN 配置"),
    ]

    for num, title, desc in steps_sr:
        doc.add_heading(f"步骤{num}: {title}", level=3)
        for line in desc.split("\n"):
            if line.strip():
                doc.add_paragraph(line, style='List Bullet')

    doc.add_page_break()

    # =========================================
    # SECTION 4: Clash Verge (Windows / macOS)
    # =========================================
    doc.add_heading(
        "四、Clash Verge（Windows / macOS）配置教程",
        level=1
    )

    doc.add_paragraph(
        "Clash Verge 是 Windows/macOS 平台最推荐的 Clash Meta 客户端，"
        "支持丰富的分流规则和 TUN 模式。"
    )

    steps_clash = [
        ("1", "下载安装",
         "1. 从 GitHub 下载对应系统版本:\n"
         "   https://github.com/clash-verge-rev/clash-verge-rev/releases\n"
         "2. Windows 下载 .exe，macOS 下载 .dmg 并安装"),
        ("2", "导入订阅",
         "1. 打开 Clash Verge\n"
         "2. 「配置」标签 → 「新建」\n"
         "3. 类型选「Remote」\n"
         "4. URL 填入:\n"
         f"   {CLASH_URL}\n"
         "5. 名称填「公司VPN」\n"
         "6. 点击「确定」"),
        ("3", "更新与选择节点",
         "1. 点击「更新」按钮拉取节点\n"
         "2. 切换到「代理」标签\n"
         "3. 选择节点\n"
         "4. 规则模式下常用「自动选择」或手动指定地区"),
        ("4", "开启代理",
         "1. 开启「系统代理」开关（浏览器生效）\n"
         "2. 或开启「TUN 模式」（所有流量生效，需管理员权限）"),
    ]

    for num, title, desc in steps_clash:
        doc.add_heading(f"步骤{num}: {title}", level=3)
        for line in desc.split("\n"):
            if line.strip():
                doc.add_paragraph(line, style='List Bullet')

    doc.add_page_break()

    # =========================================
    # SECTION 5: FAQ
    # =========================================
    doc.add_heading("五、常见问题", level=1)

    faqs = [
        ("Q: 订阅更新失败怎么办？",
         "A: 检查网络连接 → 确保能访问 sub.tamovpn.top → "
         "等待5分钟后重试。订阅每6小时自动更新一次，"
         "如长时间失败请联系 IT 管理员。"),
        ("Q: 所有节点都无法连接？",
         "A: 切换到不同的协议节点（如从 trojan 切到 vless） → "
         "尝试不同地区节点（日本/新加坡/美国） → "
         "如全部失败可能是本地网络限制了代理协议，"
         "请联系 IT 管理员配置备用方案。"),
        ("Q: 速度慢怎么办？",
         "A: 在客户端中查看节点延迟 → 选择延迟最低（<200ms）的节点 → "
         "Clash 用户可在「代理」页面测试节点延迟。"
         "优先选择地理距离近的地区（香港/日本/新加坡）。"),
        ("Q: 苹果手机怎么下载 Shadowrocket？",
         "A: 需要非中国区 Apple ID（美区/港区/台区均可）。"
         "如没有外区账号，请联系 IT 管理员协助。"),
        ("Q: 国内网络无法访问 GitHub 下载客户端？",
         "A: 可使用国内镜像下载：\n"
         "  · GitHub 镜像站: https://ghproxy.com/\n"
         "  · 或直接使用本手册附录中的「国内镜像」地址下载"),
        ("Q: 订阅地址在国内无法访问？",
         "A: 我们已配置双线路访问:\n"
         "  · 主线路(Cloudflare): https://sub.tamovpn.top/sub.txt?token=9eeeb06655487c286527001b\n"
         "  · 备用线路(CDN镜像): 见附录「国内镜像订阅地址」\n"
         "  优先使用主线路，若不可用则切换备用线路。"),
        ("Q: Token 泄露了怎么办？",
         "A: 立即联系 IT 管理员更换 Token。旧 Token 会在更换后立即失效。"),
    ]

    for q, a in faqs:
        p_q = doc.add_paragraph()
        run = p_q.add_run(q)
        run.bold = True
        run.font.size = Pt(11)
        p_a = doc.add_paragraph(a)
        p_a.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # =========================================
    # SECTION 6: Appendix
    # =========================================
    doc.add_heading("六、附录：客户端下载链接", level=1)

    add_styled_table(
        doc,
        ["客户端", "平台", "官方下载地址"],
        [
            ["v2rayNG", "Android",
             "https://github.com/2dust/v2rayNG/releases"],
            ["Clash Verge", "Windows/macOS",
             "https://github.com/clash-verge-rev/clash-verge-rev/releases"],
            ["Shadowrocket", "iOS",
             "App Store 外区账号搜索下载"],
            ["Quantumult X", "iOS",
             "App Store 外区账号搜索下载"],
            ["Sing-Box", "全平台",
             "https://github.com/SagerNet/sing-box/releases"],
        ]
    )

    doc.add_heading("国内镜像下载（GitHub 加速）", level=2)

    doc.add_paragraph(
        "如果无法直接访问 GitHub 下载客户端，可使用以下镜像地址。"
        f"将官方 Release 链接前的 https://github.com 替换为 {GHPROXY} 即可。"
    )

    add_styled_table(
        doc,
        ["客户端", "平台", "国内镜像地址"],
        [
            ["v2rayNG", "Android",
             f"{GHPROXY}/2dust/v2rayNG/releases"],
            ["Clash Verge", "Windows/macOS",
             f"{GHPROXY}/clash-verge-rev/clash-verge-rev/releases"],
            ["Sing-Box", "全平台",
             f"{GHPROXY}/SagerNet/sing-box/releases"],
        ]
    )

    doc.add_heading("订阅地址汇总", level=2)

    add_styled_table(
        doc,
        ["用途", "地址"],
        [
            ["V2Ray 订阅（主线路）", SUB_URL],
            ["Clash Meta 订阅（主线路）", CLASH_URL],
            ["Web 控制台", WEB_URL],
            ["健康检查", f"https://{HOST}/health"],
        ]
    )

    doc.add_heading("国内镜像订阅地址（备用线路）", level=2)

    doc.add_paragraph(
        "当主线路（sub.tamovpn.top）在国内无法访问时，"
        "可使用以下 CDN 镜像地址。"
        "注意：镜像地址不含 Token 验证，更新延迟最多 24 小时。"
    )

    add_styled_table(
        doc,
        ["用途", "镜像地址"],
        [
            ["V2Ray 订阅（Base64）", SUB_MIRROR_URL],
            ["Clash Meta 订阅（YAML）", CLASH_MIRROR_URL],
        ]
    )

    # Footer note
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Token: {TOKEN}  |  订阅域名: {HOST}  |  自动更新: 每6小时"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    # Save
    doc.save(str(OUTPUT))
    print(f"Manual saved: {OUTPUT}")


if __name__ == "__main__":
    build()
